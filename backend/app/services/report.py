from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches

from app.core.config import UPLOAD_DIR


def generate_report(
    image_path: str,
    result_image_path: str,
    detections: list[dict],
    model_name: str,
    inference_time_ms: float,
    use_clahe: bool,
    ai_analysis: str | None = None,
) -> str:
    doc = Document()
    doc.add_heading("漂浮物检测报告", level=1)
    doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"使用模型: {model_name}")
    doc.add_paragraph(f"CLAHE 预处理: {'启用' if use_clahe else '未启用'}")
    doc.add_paragraph(f"推理耗时: {inference_time_ms:.1f} ms")
    doc.add_paragraph(f"检测目标数: {len(detections)}")

    doc.add_heading("检测详情", level=2)
    if detections:
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "类别"
        hdr[1].text = "置信度"
        hdr[2].text = "边界框 (x1,y1,x2,y2)"
        hdr[3].text = "区域面积"
        for d in detections:
            row = table.add_row().cells
            row[0].text = d["class_name"]
            row[1].text = f"{d['confidence']:.2%}"
            row[2].text = f"({d['x1']:.0f},{d['y1']:.0f},{d['x2']:.0f},{d['y2']:.0f})"
            area = (d["x2"] - d["x1"]) * (d["y2"] - d["y1"])
            row[3].text = f"{area:.0f} px²"
    else:
        doc.add_paragraph("未检测到任何目标。")

    if ai_analysis:
        doc.add_heading("AI 分析建议", level=2)
        doc.add_paragraph(ai_analysis)

    if Path(result_image_path).exists():
        doc.add_heading("检测结果图", level=2)
        doc.add_picture(result_image_path, width=Inches(5))

    output_dir = UPLOAD_DIR / "reports"
    output_dir.mkdir(parents=True, exist_ok=True)
    filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    output_path = output_dir / filename
    doc.save(str(output_path))
    return str(output_path)
